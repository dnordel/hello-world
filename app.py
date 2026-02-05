"""
Structured Preschool Interview Tool (Offline)

This module contains:
- JSON rubric loading + validation
- Disqualifier signal library loading (optional JSON)
- Scoring engine (weighted totals + threshold logic + override locks)
- Draft saving/loading (JSON)
- Docx export (python-docx)
- Tkinter GUI app skeleton and key screens

Notes:
- This code assumes `rubric.json` exists next to this file.
- `disqualifier_signals.json` is optional.
- The GUI portions shown here are based on your pasted snippet; some screens
  (Settings, Open Draft, etc.) were truncated in your paste, so placeholders
  are included where needed.
"""

from __future__ import annotations

import json
import re
import traceback
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Optional

import tkinter as tk
from tkinter import BooleanVar, END, IntVar, StringVar, filedialog, messagebox, ttk
from tkinter import font as tkfont

from docx import Document


# =========================
# App constants and defaults
# =========================

APP_TITLE = "Structured Preschool Interview Tool"

# Directory containing this script file.
APP_DIR = Path(__file__).resolve().parent

# Default data files. rubric.json is required; signals file is optional.
DEFAULT_RUBRIC_PATH = APP_DIR / "rubric.json"
DEFAULT_SIGNALS_PATH = APP_DIR / "disqualifier_signals.json"

# Default output base directory:
# - drafts go under <base>/drafts
# - final reports go under <base>/final
DEFAULT_BASE_DIR = APP_DIR / "interviews"

# Font sizing controls for accessibility.
DEFAULT_FONT_SIZE = 10
MIN_FONT_SIZE = 8
MAX_FONT_SIZE = 18

DEFAULT_SCHOOL_OPTIONS = [
    "Hawthorne",
    "Palmdale",
    "North Long Beach",
]


# =========================
# Utility helpers
# =========================

def sanitize_filename(name: str) -> str:
    """
    Replace characters that are illegal or awkward in file names,
    and collapse whitespace.
    """
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name or "Unknown"


def is_valid_date_yyyy_mm_dd(value: str) -> bool:
    """
    Validate date string in YYYY-MM-DD format.
    """
    try:
        datetime.strptime(value, "%Y-%m-%d")
        return True
    except ValueError:
        return False


# =========================
# Data loading and validation
# =========================

class RubricLoader:
    """
    Loads rubric.json and validates required structure.

    rubric.json is expected to include:
    - metadata
    - scoring
    - tracks
    - traits
    - absolute_disqualifiers
    """

    def __init__(self, rubric_path: Path):
        self.rubric_path = Path(rubric_path)
        self.data = self._load()

    def _load(self) -> dict[str, Any]:
        """
        Read and parse rubric JSON, then validate.
        """
        if not self.rubric_path.exists():
            raise FileNotFoundError(f"Rubric file not found: {self.rubric_path}")

        with self.rubric_path.open("r", encoding="utf-8") as f:
            data = json.load(f)

        self.validate(data)
        return data

    @staticmethod
    def validate(data: dict[str, Any]) -> None:
        """
        Enforce a minimum schema so the GUI and scoring don't crash mid-interview.
        """
        required = ["metadata", "scoring", "tracks", "traits", "absolute_disqualifiers"]
        for key in required:
            if key not in data:
                raise ValueError(f"rubric.json missing required key: {key}")

        if not isinstance(data["traits"], list) or not data["traits"]:
            raise ValueError("rubric.json requires non-empty list: traits")

        # Trait schema validation: only checks presence of keys (not content correctness).
        for trait in data["traits"]:
            required_trait_keys = [
                "id",
                "name",
                "priority",
                "weight",
                "primary_question",
                "descriptors",
                "sample_answers",
                "applicable_tracks",
            ]
            for k in required_trait_keys:
                if k not in trait:
                    raise ValueError(f"Trait missing '{k}': {trait}")

    def get_traits_for_track(self, track_key: str) -> list[dict[str, Any]]:
        """
        Return traits applicable to a given track.
        A trait applies if 'all' is included, or the track_key is included.
        """
        traits: list[dict[str, Any]] = []
        for t in self.data["traits"]:
            applicable = t.get("applicable_tracks", [])
            if "all" in applicable or track_key in applicable:
                traits.append(t)
        return traits


class DisqualifierSignalLibrary:
    """
    Loads disqualifier_signals.json (optional) and indexes it by trait_id.

    File shape expected (loosely):
    {
      "questions": [
        {
          "trait_id": "trait_1" or "1" or "some_trait_key",
          "question_id": "...",
          "primary_question": "...",
          "disqualifier_signals": [
             {
               "disqualifier_type": "emotion_dysregulation",
               "auto_disqualify_if_confirmed": true,
               "examples": [...],
               "probe_to_confirm": "..."
             }
          ]
        }
      ]
    }
    """

    def __init__(self, path: Path):
        self.path = Path(path)
        self.data = self._load()
        self.by_trait_id = self._build_index()

    def _load(self) -> dict[str, Any]:
        """
        Load the JSON file if present; otherwise return an empty library.
        """
        if not self.path.exists():
            return {"questions": []}

        with self.path.open("r", encoding="utf-8") as f:
            return json.load(f)

    def _build_index(self) -> dict[str, dict[str, Any]]:
        """
        Index questions by trait id, normalizing numeric trait ids to "trait_<n>".
        """
        out: dict[str, dict[str, Any]] = {}
        for q in self.data.get("questions", []):
            raw_trait = str(q.get("trait_id", "")).strip()
            if raw_trait.isdigit():
                out[f"trait_{raw_trait}"] = q
            elif raw_trait:
                out[raw_trait] = q
        return out

    def get_for_trait(self, trait_id: str) -> Optional[dict[str, Any]]:
        """
        Return the signal block for a specific trait_id, if configured.
        """
        return self.by_trait_id.get(trait_id)


# =========================
# Scoring and decisions
# =========================

class ScoringEngine:
    """
    Computes:
    - weighted totals
    - percent of max
    - critical trait override flags
    - absolute disqualifier lock
    - final outcome string: "Hire" / "Borderline" / "No Hire"
    """

    @staticmethod
    def evaluate(rubric: dict[str, Any], track_key: str, trait_results: dict[str, dict[str, Any]]) -> dict[str, Any]:
        """
        Evaluate a set of per-trait inputs against the rubric scoring rules.
        """
        # Only score traits that apply to this track.
        traits = [
            t for t in rubric["traits"]
            if "all" in t["applicable_tracks"] or track_key in t["applicable_tracks"]
        ]

        rows: list[dict[str, Any]] = []
        weighted_total = 0

        # Override flags.
        critical_eq_1 = False  # Any critical trait raw score == 1
        critical_lt_3 = False  # Any critical trait raw score < 3
        disqualifier_present = False  # Any trait has absolute disqualifier checked

        for trait in traits:
            tid = trait["id"]
            state = trait_results.get(tid, {})

            # raw_score might be None until interviewer selects one.
            raw = int(state.get("raw_score", 0) or 0)

            # Weight is defined on the trait itself.
            weight = int(trait["weight"])
            weighted = raw * weight
            weighted_total += weighted

            # Absolute disqualifier is a manual checkbox in the UI.
            dq = bool(state.get("absolute_disqualifier", False))
            if dq:
                disqualifier_present = True

            is_critical = str(trait["priority"]).lower() == "critical"
            if is_critical and raw == 1:
                critical_eq_1 = True
            if is_critical and raw < 3:
                critical_lt_3 = True

            rows.append(
                {
                    "trait_id": tid,
                    "trait_name": trait["name"],
                    "priority": trait["priority"],
                    "weight": weight,
                    "raw_score": raw,
                    "weighted_score": weighted,
                    "question_notes": state.get("question_notes", ""),
                    "trait_notes": state.get("trait_notes", ""),
                    "verbatim_notes": state.get("verbatim_notes", ""),
                    "absolute_disqualifier": dq,
                    "primary_question": trait["primary_question"],
                }
            )

        # Max weighted total is stored per track.
        max_weighted = int(rubric["tracks"][track_key]["max_weighted_total"])
        pct = (weighted_total / max_weighted) * 100 if max_weighted else 0.0

        # Locked rule indicates that the outcome is forced regardless of thresholds.
        locked_rule: Optional[str] = None
        if disqualifier_present:
            locked_rule = "Any Absolute Disqualifier observed => Immediate NO HIRE"
        if critical_eq_1:
            locked_rule = "Any Critical trait raw score = 1 => Immediate NO HIRE"

        # Determine outcome.
        if disqualifier_present or critical_eq_1:
            outcome = "No Hire"
        else:
            # Keep thresholds exactly as your logic:
            # - Hire: >=80% and no critical <3
            # - Borderline: 65-79% and not critical=1 (already excluded above)
            # - Else: No Hire
            if pct >= 80 and not critical_lt_3:
                outcome = "Hire"
            elif 65 <= pct <= 79 and not critical_eq_1:
                outcome = "Borderline"
            else:
                outcome = "No Hire"

        return {
            "rows": rows,
            "weighted_total": weighted_total,
            "max_weighted_total": max_weighted,
            "percent_of_max": round(pct, 2),
            "critical_eq_1": critical_eq_1,
            "critical_lt_3": critical_lt_3,
            "disqualifier_present": disqualifier_present,
            "locked_rule": locked_rule,
            "outcome": outcome,
        }


# =========================
# Draft persistence
# =========================

class DraftManager:
    """
    Saves and loads interview drafts as JSON.
    Drafts are timestamped and stored under <base>/drafts.
    """

    def __init__(self, base_dir: Path):
        self.base_dir = Path(base_dir)
        self.drafts_dir = self.base_dir / "drafts"
        self.final_dir = self.base_dir / "final"
        self.drafts_dir.mkdir(parents=True, exist_ok=True)
        self.final_dir.mkdir(parents=True, exist_ok=True)

    def save_draft(self, payload: dict[str, Any]) -> Path:
        """
        Save a JSON draft, naming it with timestamp and candidate name.
        """
        candidate = payload.get("candidate", {}).get("name", "Unknown")
        safe = sanitize_filename(candidate or "Unknown")
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        path = self.drafts_dir / f"draft-{stamp}-{safe}.json"
        with path.open("w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
        return path

    def load_draft(self, path: Path) -> dict[str, Any]:
        """
        Load a JSON draft payload from disk.
        """
        with Path(path).open("r", encoding="utf-8") as f:
            return json.load(f)


# =========================
# DOCX report export
# =========================

class DocxExporter:
    """
    Exports a finalized interview report to a single .docx file (one per candidate).
    """

    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(self, rubric: dict[str, Any], payload: dict[str, Any], scoring: dict[str, Any]) -> Path:
        """
        Create a Word document with:
        - Candidate info
        - Score summary table
        - Override summary
        - Trait-by-trait notes
        - Global disqualifiers
        - Observed disqualifier evidence lines
        """
        candidate = payload["candidate"]
        cname = candidate["name"]
        interview_date = candidate["interview_date"]
        track_key = candidate["track"]
        school = candidate.get("school", "")
        track_label = rubric["tracks"][track_key]["label"]

        doc = Document()
        doc.add_heading("Structured Behavioral Interview Report", level=1)

        # Candidate meta section.
        doc.add_paragraph(f"Candidate Name: {cname}")
        doc.add_paragraph(f"Interview Date: {interview_date}")
        doc.add_paragraph(f"School/Location: {school}")
        doc.add_paragraph(f"Track: {track_label}")

        # Score summary.
        doc.add_heading("Score Summary", level=2)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Trait"
        hdr[1].text = "Priority"
        hdr[2].text = "Weight"
        hdr[3].text = "Raw Score"
        hdr[4].text = "Weighted Score"

        for row in scoring["rows"]:
            cells = table.add_row().cells
            cells[0].text = row["trait_name"]
            cells[1].text = row["priority"]
            cells[2].text = str(row["weight"])
            cells[3].text = str(row["raw_score"])
            cells[4].text = str(row["weighted_score"])

        doc.add_paragraph(f"Weighted Total: {scoring['weighted_total']} / {scoring['max_weighted_total']}")
        doc.add_paragraph(f"Percent of Max: {scoring['percent_of_max']}%")
        doc.add_paragraph(f"Final Outcome: {scoring['outcome']}")

        # Overrides.
        doc.add_heading("Override Summary", level=2)
        doc.add_paragraph(f"Any Critical trait = 1: {'Yes' if scoring['critical_eq_1'] else 'No'}")
        doc.add_paragraph(f"Any Absolute Disqualifier observed: {'Yes' if scoring['disqualifier_present'] else 'No'}")
        doc.add_paragraph(f"Outcome lock rule: {scoring['locked_rule'] if scoring['locked_rule'] else 'None'}")

        # Trait detail.
        doc.add_heading("Trait-by-Trait Detail", level=2)
        for idx, row in enumerate(scoring["rows"], start=1):
            doc.add_heading(f"{idx}. {row['trait_name']}", level=3)
            doc.add_paragraph(f"Priority: {row['priority']} | Weight: x{row['weight']}")
            doc.add_paragraph(f"Primary Question: {row['primary_question']}")
            doc.add_paragraph(f"Selected Raw Score: {row['raw_score']}")
            doc.add_paragraph(f"Question Notes: {row['question_notes']}")
            doc.add_paragraph(f"Trait Notes: {row['trait_notes']}")
            doc.add_paragraph(f"Verbatim quote/notes: {row['verbatim_notes']}")

        # Global disqualifiers.
        doc.add_heading("Global Disqualifiers", level=2)
        for d in rubric["absolute_disqualifiers"]:
            doc.add_paragraph(f"- {d}")

        # Evidence section.
        doc.add_paragraph("Observed disqualifier evidence (from verbatim notes):")
        evidence_added = False
        for row in scoring["rows"]:
            if row["absolute_disqualifier"] and row["verbatim_notes"].strip():
                doc.add_paragraph(f"- {row['trait_name']}: {row['verbatim_notes'].strip()}")
                evidence_added = True
        if not evidence_added:
            doc.add_paragraph("- None recorded")

        # Output filename format.
        school_part = sanitize_filename(school) if school else "UnknownSchool"
        filename = f"{interview_date} - {school_part} - {sanitize_filename(cname)} - Interview.docx"
        out_path = self.output_dir / filename
        doc.save(out_path)
        return out_path


# =========================
# Interview state container
# =========================

@dataclass
class InterviewState:
    """
    In-memory state for the current interview session.

    trait_inputs schema:
    {
      trait_id: {
        "raw_score": int|None,
        "question_notes": str,
        "trait_notes": str,
        "verbatim_notes": str,
        "absolute_disqualifier": bool
      }
    }
    """

    candidate_name: str = ""
    interview_date: str = ""
    school: str = ""
    track: str = ""
    current_index: int = 0

    # Use default_factory to avoid shared mutable defaults between instances.
    trait_inputs: dict[str, dict[str, Any]] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        """
        Convert state to a JSON-friendly payload.
        """
        return {
            "candidate": {
                "name": self.candidate_name,
                "interview_date": self.interview_date,
                "school": self.school,
                "track": self.track,
            },
            "current_index": self.current_index,
            "trait_inputs": self.trait_inputs,
        }


# =========================
# Tkinter GUI app
# =========================

class InterviewApp(tk.Tk):
    """
    Main Tkinter application.

    Key responsibilities:
    - Load rubric and signals
    - Maintain InterviewState
    - Render screens (start, candidate info, trait screens)
    - Persist drafts
    - Validate and finalize to DOCX
    """

    def __init__(self):
        super().__init__()

        # Basic window config.
        self.title(APP_TITLE)
        self.geometry("1100x800")

        # Persistable user settings (in your snippet these were in-memory only).
        self.settings: dict[str, Any] = {
            "base_dir": str(DEFAULT_BASE_DIR),
            "font_size": DEFAULT_FONT_SIZE,
        }

        self.school_options = DEFAULT_SCHOOL_OPTIONS.copy()

        # Load rubric and signals.
        self.rubric_loader = RubricLoader(DEFAULT_RUBRIC_PATH)
        self.rubric = self.rubric_loader.data
        self.signals = DisqualifierSignalLibrary(DEFAULT_SIGNALS_PATH)

        # Initialize state with today's date.
        self.state = InterviewState(interview_date=date.today().isoformat())

        # Active traits list is set after candidate selects track.
        self.active_traits: list[dict[str, Any]] = []

        # Theme must be configured before widgets are created.
        self._configure_theme()
        self.apply_font_size(self.settings["font_size"])

        # Build layout skeleton.
        self._build_layout()

        # Show first screen.
        self.show_start_screen()

    # -------------------------
    # Layout and theming
    # -------------------------

    def _build_layout(self) -> None:
        """
        Build the fixed outer layout:
        - top toolbar
        - scrollable main page area
        - footer action bar
        """
        self.toolbar = ttk.Frame(self)
        self.toolbar.pack(fill="x", padx=8, pady=4)

        ttk.Label(self.toolbar, text="Text Size:").pack(side="left")
        ttk.Button(self.toolbar, text="A-", command=lambda: self.adjust_font_size(-1)).pack(side="left", padx=2)

        self.font_label = ttk.Label(self.toolbar, text=str(self.settings["font_size"]))
        self.font_label.pack(side="left", padx=2)

        ttk.Button(self.toolbar, text="A+", command=lambda: self.adjust_font_size(1)).pack(side="left", padx=2)

        # Main scroll container.
        self.main_holder = ttk.Frame(self)
        self.main_holder.pack(fill="both", expand=True)

        # Canvas enables a scrollable "page".
        self.canvas = tk.Canvas(self.main_holder, highlightthickness=0, bg="#f3f5f8")
        self.v_scroll = ttk.Scrollbar(self.main_holder, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.v_scroll.set)

        self.v_scroll.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        # This frame is the actual content area embedded in the canvas.
        self.page_frame = ttk.Frame(self.canvas)
        self.page_window = self.canvas.create_window((0, 0), window=self.page_frame, anchor="nw")

        # Keep scroll region synced with content size.
        self.page_frame.bind("<Configure>", self._on_frame_configure)
        self.canvas.bind("<Configure>", self._on_canvas_configure)

        # Mouse wheel scrolling for Windows and Linux.
        self.bind_all("<MouseWheel>", self._on_mousewheel)
        self.bind_all("<Button-4>", self._on_mousewheel)
        self.bind_all("<Button-5>", self._on_mousewheel)

        # Footer bar for navigation buttons.
        self.footer_separator = ttk.Separator(self, orient="horizontal")
        self.footer_separator.pack(fill="x")

        self.footer = ttk.Frame(self, padding=(8, 6))
        self.footer.pack(fill="x")

    def _configure_theme(self) -> None:
        """
        Apply a light theme and baseline ttk styles.
        """
        self.configure(background="#f3f5f8")
        style = ttk.Style(self)

        # "clam" is widely available and looks modern. If unavailable, fallback silently.
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        # Global backgrounds.
        style.configure("TFrame", background="#f3f5f8")
        style.configure("TLabel", background="#f3f5f8")

        # Group boxes with white background for contrast.
        style.configure("TLabelframe", background="#ffffff", borderwidth=1, relief="solid")
        style.configure(
            "TLabelframe.Label",
            background="#ffffff",
            foreground="#1f2937",
            font=("TkDefaultFont", self.settings["font_size"] + 1, "bold"),
        )

        # Buttons and inputs.
        style.configure("TButton", padding=(10, 6))
        style.configure("TRadiobutton", background="#f3f5f8")
        style.configure("TCheckbutton", background="#f3f5f8")
        style.configure("TCombobox", padding=4)

    def _on_frame_configure(self, _event: tk.Event) -> None:
        """
        Whenever the page_frame changes size, update the scrollable region.
        """
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_canvas_configure(self, event: tk.Event) -> None:
        """
        Make the embedded page_frame match the canvas width so content reflows.
        """
        self.canvas.itemconfig(self.page_window, width=event.width)

    def _on_mousewheel(self, event: tk.Event) -> None:
        """
        Cross-platform mouse wheel handler:
        - Windows uses event.delta
        - Some Linux setups use Button-4/Button-5
        """
        if getattr(event, "num", None) == 4:
            self.canvas.yview_scroll(-3, "units")
        elif getattr(event, "num", None) == 5:
            self.canvas.yview_scroll(3, "units")
        else:
            # event.delta is typically +/-120 per notch on Windows
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    # -------------------------
    # Font sizing utilities
    # -------------------------

    def apply_font_size(self, size: int) -> None:
        """
        Apply a global UI font size for accessibility.
        """
        size = max(MIN_FONT_SIZE, min(MAX_FONT_SIZE, int(size)))
        self.settings["font_size"] = size

        # Reconfigure theme to ensure style fonts update with new size.
        self._configure_theme()

        # Update Tk named fonts so labels, entries, etc. scale.
        for font_name in ("TkDefaultFont", "TkTextFont", "TkMenuFont", "TkHeadingFont", "TkCaptionFont"):
            try:
                f = tkfont.nametofont(font_name)
                f.configure(size=size)
            except tk.TclError:
                # Some named fonts may not exist on all platforms.
                pass

        if hasattr(self, "font_label"):
            self.font_label.config(text=str(size))

    def adjust_font_size(self, delta: int) -> None:
        """
        Increase/decrease global font size in one step.
        """
        self.apply_font_size(self.settings["font_size"] + delta)

    # -------------------------
    # Page utilities
    # -------------------------

    def scroll_top(self) -> None:
        """Scroll to the top of the canvas."""
        self.canvas.yview_moveto(0)

    def clear_page(self) -> None:
        """
        Remove all widgets from the page area and reset scroll position.
        Also clears footer actions.
        """
        for child in self.page_frame.winfo_children():
            child.destroy()
        self.clear_footer()
        self.scroll_top()

    def clear_footer(self) -> None:
        """Remove all footer buttons."""
        for child in self.footer.winfo_children():
            child.destroy()

    def set_footer_actions(self, left_actions=None, right_actions=None) -> None:
        """
        Render footer buttons in consistent positions.
        left_actions and right_actions are lists of (label, callback).
        """
        self.clear_footer()

        left = ttk.Frame(self.footer)
        left.pack(side="left")
        for label, command in (left_actions or []):
            ttk.Button(left, text=label, command=command).pack(side="left", padx=4)

        right = ttk.Frame(self.footer)
        right.pack(side="right")
        for label, command in (right_actions or []):
            ttk.Button(right, text=label, command=command).pack(side="right", padx=4)

    # -------------------------
    # Screens
    # -------------------------

    def show_start_screen(self) -> None:
        """
        Start screen: New Interview, Open Draft, Settings, Exit.
        """
        self.clear_page()

        frm = ttk.Frame(self.page_frame, padding=20)
        frm.pack(fill="both", expand=True)

        ttk.Label(frm, text=APP_TITLE, font=("TkDefaultFont", self.settings["font_size"] + 6, "bold")).pack(pady=10)
        ttk.Button(frm, text="New Interview", command=self.new_interview).pack(fill="x", pady=5)
        ttk.Button(frm, text="Open Draft", command=self.open_draft).pack(fill="x", pady=5)
        ttk.Button(frm, text="Settings", command=self.open_settings).pack(fill="x", pady=5)
        ttk.Button(frm, text="Exit", command=self.destroy).pack(fill="x", pady=5)

        self.set_footer_actions()

    def new_interview(self) -> None:
        """
        Reset state and go to candidate info.
        """
        self.state = InterviewState(interview_date=date.today().isoformat())
        self.active_traits = []
        self.show_candidate_info()

    def open_draft(self) -> None:
        """
        Load a previously saved draft JSON and resume.
        """
        base_dir = Path(self.settings["base_dir"])
        dm = DraftManager(base_dir)

        path = filedialog.askopenfilename(
            title="Open Draft",
            initialdir=str(dm.drafts_dir),
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            payload = dm.load_draft(Path(path))
            cand = payload.get("candidate", {})
            self.state.candidate_name = cand.get("name", "")
            self.state.interview_date = cand.get("interview_date", date.today().isoformat())
            self.state.school = cand.get("school", "")
            self.state.track = cand.get("track", "")
            self.state.current_index = int(payload.get("current_index", 0) or 0)
            self.state.trait_inputs = payload.get("trait_inputs", {}) or {}

            # Recompute active traits from the selected track so navigation works.
            if self.state.track:
                self.active_traits = self.rubric_loader.get_traits_for_track(self.state.track)
                start_idx = max(0, self.state.current_index - 1)
                self.show_trait_screen(start_idx)
            else:
                self.show_candidate_info()

        except Exception as exc:
            messagebox.showerror("Open Draft Error", f"{exc}\n\n{traceback.format_exc()}")

    def open_settings(self) -> None:
        """
        Simple settings window:
        - Base output folder
        - Font size slider
        """
        top = tk.Toplevel(self)
        top.title("Settings")
        top.geometry("650x260")

        path_var = StringVar(value=self.settings["base_dir"])
        size_var = IntVar(value=self.settings["font_size"])

        row = ttk.Frame(top, padding=10)
        row.pack(fill="x")
        ttk.Label(row, text="Base output folder:").pack(anchor="w")
        ttk.Entry(row, textvariable=path_var).pack(fill="x", pady=5)

        def browse():
            d = filedialog.askdirectory(title="Select Base Output Folder", initialdir=path_var.get() or str(DEFAULT_BASE_DIR))
            if d:
                path_var.set(d)

        ttk.Button(row, text="Browse...", command=browse).pack(anchor="w")

        font_row = ttk.Frame(top, padding=10)
        font_row.pack(fill="x")
        ttk.Label(font_row, text="Font size:").pack(anchor="w")

        slider = ttk.Scale(
            font_row,
            from_=MIN_FONT_SIZE,
            to=MAX_FONT_SIZE,
            orient="horizontal",
            variable=size_var,
        )
        slider.pack(fill="x", pady=8)

        btn_row = ttk.Frame(top, padding=10)
        btn_row.pack(fill="x")

        def save():
            self.settings["base_dir"] = path_var.get().strip() or str(DEFAULT_BASE_DIR)
            self.apply_font_size(int(size_var.get()))
            top.destroy()

        ttk.Button(btn_row, text="Save", command=save).pack(side="right")
        ttk.Button(btn_row, text="Cancel", command=top.destroy).pack(side="right", padx=6)

    def show_candidate_info(self) -> None:
        """
        Candidate info screen:
        - candidate name
        - interview date
        - school selection or add
        - track selection
        """
        self.clear_page()

        frm = ttk.Frame(self.page_frame, padding=20)
        frm.pack(fill="both", expand=True)

        ttk.Label(frm, text="Step 1: Candidate Info", font=("TkDefaultFont", self.settings["font_size"] + 4, "bold")).pack(anchor="w", pady=8)

        name_var = StringVar(value=self.state.candidate_name)
        date_var = StringVar(value=self.state.interview_date or date.today().isoformat())
        school_var = StringVar(value=self.state.school)
        track_var = StringVar(value=self.state.track)

        ttk.Label(frm, text="Candidate Name (required)").pack(anchor="w")
        ttk.Entry(frm, textvariable=name_var).pack(fill="x", pady=4)

        ttk.Label(frm, text="Interview Date YYYY-MM-DD (required)").pack(anchor="w")
        ttk.Entry(frm, textvariable=date_var).pack(fill="x", pady=4)

        ttk.Label(frm, text="School (required)").pack(anchor="w")
        school_row = ttk.Frame(frm)
        school_row.pack(fill="x", pady=4)

        school_combo = ttk.Combobox(school_row, textvariable=school_var, values=self.school_options)
        school_combo.pack(side="left", fill="x", expand=True)

        # Make combobox behave like a dropdown on click.
        def open_school_dropdown(_event=None):
            school_combo.focus_set()
            school_combo.after_idle(lambda: school_combo.event_generate("<Down>"))

        school_combo.bind("<Button-1>", open_school_dropdown)

        def add_school():
            value = school_var.get().strip()
            if not value:
                messagebox.showerror("Validation", "Enter a school name before adding.")
                return
            if value not in self.school_options:
                self.school_options.append(value)
            school_combo.configure(values=self.school_options)
            school_var.set(value)

        ttk.Button(school_row, text="Add School", command=add_school).pack(side="left", padx=6)

        ttk.Label(frm, text="Track (required)").pack(anchor="w", pady=(10, 0))

        tracks = [(k, self.rubric["tracks"][k]["label"]) for k in self.rubric["tracks"].keys()]
        for k, label in tracks:
            ttk.Radiobutton(frm, text=label, variable=track_var, value=k).pack(anchor="w")

        threshold_lbl = ttk.Label(frm, text="")
        threshold_lbl.pack(anchor="w", pady=10)

        def refresh_thresholds(*_):
            t = track_var.get()
            if t and t in self.rubric["tracks"]:
                cfg = self.rubric["tracks"][t]
                threshold_lbl.config(
                    text=(
                        f"Thresholds: Hire >= {cfg.get('hire_percent', '??')}% "
                        f"| Borderline {cfg.get('borderline_min_percent', '??')}%-{cfg.get('borderline_max_percent', '??')}% "
                        f"| No Hire < {cfg.get('borderline_min_percent', '??')}% "
                        f"or any Critical < 3 "
                        f"or any Critical = 1 "
                        f"or any Absolute Disqualifier"
                    )
                )
            else:
                threshold_lbl.config(text="")

        track_var.trace_add("write", refresh_thresholds)
        refresh_thresholds()

        def go_next():
            """
            Validate candidate inputs, initialize per-trait state,
            and move to the first trait screen.
            """
            name = name_var.get().strip()
            iv_date = date_var.get().strip()
            school = school_var.get().strip()
            track = track_var.get().strip()

            if not name:
                messagebox.showerror("Validation", "Candidate Name is required.")
                return
            if not is_valid_date_yyyy_mm_dd(iv_date):
                messagebox.showerror("Validation", "Interview Date must be valid YYYY-MM-DD.")
                return
            if not school:
                messagebox.showerror("Validation", "School selection is required.")
                return
            if not track:
                messagebox.showerror("Validation", "Track selection is required.")
                return

            # Save candidate info.
            self.state.candidate_name = name
            self.state.interview_date = iv_date
            self.state.school = school
            self.state.track = track

            # Build active trait list and ensure defaults exist for each trait.
            self.active_traits = self.rubric_loader.get_traits_for_track(track)
            for trait in self.active_traits:
                self.state.trait_inputs.setdefault(
                    trait["id"],
                    {
                        "raw_score": None,
                        "question_notes": "",
                        "trait_notes": "",
                        "verbatim_notes": "",
                        "absolute_disqualifier": False,
                    },
                )

            # Index is 1-based in your saved drafts; screen index is 0-based.
            self.state.current_index = 1
            self.show_trait_screen(0)

        self.set_footer_actions(
            left_actions=[("Back to Start", self.show_start_screen)],
            right_actions=[("Next", go_next)],
        )

    def show_disqualifier_reference(self) -> None:
        """
        Popup window listing global absolute disqualifiers from rubric.json.
        """
        top = tk.Toplevel(self)
        top.title("Absolute Disqualifiers")
        top.geometry("760x360")

        text = tk.Text(top, wrap="word")
        text.pack(fill="both", expand=True)

        text.insert(END, "Absolute Disqualifiers (Global)\n\n")
        for item in self.rubric["absolute_disqualifiers"]:
            text.insert(END, f"- {item}\n")

        text.config(state="disabled")

    def _render_signal_examples(self, parent: ttk.Frame, trait_id: str) -> None:
        """
        Render disqualifier signal examples for the given trait into a styled Text widget.
        """
        data = self.signals.get_for_trait(trait_id)

        box = ttk.LabelFrame(parent, text="Disqualifier Signal Examples")
        box.pack(fill="both", pady=8, expand=True)

        ttk.Label(
            box,
            text="Use these as probe prompts; they are examples to help pattern-match risk signals.",
            wraplength=1020,
            foreground="#334155",
        ).pack(anchor="w", padx=10, pady=(8, 0))

        text = tk.Text(
            box,
            height=18,
            wrap="word",
            relief="flat",
            bg="#f8fafc",
            padx=12,
            pady=10,
            font=("TkDefaultFont", self.settings["font_size"]),
        )

        ybar = ttk.Scrollbar(box, orient="vertical", command=text.yview)
        text.configure(yscrollcommand=ybar.set)

        ybar.pack(side="right", fill="y", pady=8)
        text.pack(side="left", fill="both", expand=True, padx=8, pady=8)

        # Text tags for readability.
        text.tag_configure("header", font=("TkDefaultFont", self.settings["font_size"] + 1, "bold"), foreground="#0f172a")
        text.tag_configure("meta", foreground="#334155")
        text.tag_configure("signal", font=("TkDefaultFont", self.settings["font_size"], "bold"), foreground="#1d4ed8")
        text.tag_configure("probe", foreground="#7c2d12")

        if not data:
            text.insert(END, "No signal examples configured for this trait.", "meta")
            text.config(state="disabled")
            return

        # Basic context.
        text.insert(END, "Question Context\n", "header")
        text.insert(END, f"Question ID: {data.get('question_id', '')}\n", "meta")
        text.insert(END, f"Primary question: {data.get('primary_question', '')}\n\n", "meta")

        # Accept either key name to make the JSON tolerant.
        signal_items = data.get("disqualifier_signals") or data.get("signals") or []
        if not signal_items:
            text.insert(END, "No signal examples configured for this trait.", "meta")
            text.config(state="disabled")
            return

        for idx, item in enumerate(signal_items, start=1):
            raw_type = item.get("disqualifier_type", "")
            friendly_type = raw_type.replace("_", " ").title() if raw_type else "Unspecified"
            auto = "Yes" if item.get("auto_disqualify_if_confirmed") else "No"

            text.insert(END, f"Signal {idx}: {friendly_type}\n", "signal")
            text.insert(END, f"Auto disqualify if confirmed: {auto}\n", "meta")

            examples = item.get("examples", [])
            if examples:
                for ex in examples:
                    text.insert(END, f"• {ex}\n")
            else:
                text.insert(END, "• No examples listed.\n", "meta")

            probe = item.get("probe_to_confirm", "")
            if probe:
                text.insert(END, f"Probe to confirm: {probe}\n", "probe")

            text.insert(END, "\n")

        text.config(state="disabled")

    def show_trait_screen(self, idx: int) -> None:
        """
        Trait screen:
        - Displays trait name, priority, weight, primary question
        - Shows descriptor ladder and sample answers (read-only)
        - Captures: raw score, question notes, trait notes, verbatim notes
        - Captures: absolute disqualifier checkbox
        - Navigation: Back/Next/Save Draft/Finalize
        """
        self.clear_page()

        if idx < 0 or idx >= len(self.active_traits):
            # Defensive fallback: go to candidate info if out of bounds.
            self.show_candidate_info()
            return

        trait = self.active_traits[idx]
        tid = trait["id"]

        # Ensure a default state exists for this trait.
        state = self.state.trait_inputs.setdefault(
            tid,
            {
                "raw_score": None,
                "question_notes": "",
                "trait_notes": "",
                "verbatim_notes": "",
                "absolute_disqualifier": False,
            },
        )

        frm = ttk.Frame(self.page_frame, padding=12)
        frm.pack(fill="both", expand=True)

        # Header block.
        intro = ttk.LabelFrame(frm, text=f"Trait {idx + 1} of {len(self.active_traits)}")
        intro.pack(fill="x", pady=6)

        ttk.Label(intro, text=trait["name"], font=("TkDefaultFont", self.settings["font_size"] + 4, "bold")).pack(anchor="w", padx=10, pady=(8, 2))
        ttk.Label(intro, text=f"Priority: {trait['priority']} | Weight: x{trait['weight']}", foreground="#334155").pack(anchor="w", padx=10)
        ttk.Label(intro, text=f"Primary Question: {trait['primary_question']}", wraplength=1050).pack(anchor="w", padx=10, pady=(4, 10))

        # Ladder.
        ladder_frame = ttk.LabelFrame(frm, text="Scoring descriptors (1-5)")
        ladder_frame.pack(fill="x", pady=4)
        for n in [5, 4, 3, 2, 1]:
            ttk.Label(ladder_frame, text=f"{n}: {trait['descriptors'][str(n)]}", wraplength=1030).pack(anchor="w")

        # Sample answers.
        sample_frame = ttk.LabelFrame(frm, text="Sample answers (display only)")
        sample_frame.pack(fill="x", pady=4)
        for n in [5, 4, 3, 2, 1]:
            ttk.Label(sample_frame, text=f"{n}: {trait['sample_answers'][str(n)]}", wraplength=1030).pack(anchor="w")

        # Optional disqualifier signal reference.
        self._render_signal_examples(frm, tid)

        # Variables bound to UI.
        raw_var = IntVar(value=int(state["raw_score"]) if state.get("raw_score") else 0)
        dq_var = BooleanVar(value=bool(state.get("absolute_disqualifier", False)))

        # Raw score selector.
        score_row = ttk.Frame(frm)
        score_row.pack(fill="x", pady=6)

        ttk.Label(score_row, text="Raw score (required):").pack(side="left")
        for n in [1, 2, 3, 4, 5]:
            ttk.Radiobutton(score_row, text=str(n), value=n, variable=raw_var).pack(side="left")

        # Disqualifier checkbox and reference button.
        dq_row = ttk.Frame(frm)
        dq_row.pack(fill="x", pady=4)

        ttk.Checkbutton(dq_row, text="Absolute disqualifier observed (for this trait)", variable=dq_var).pack(side="left")
        ttk.Button(dq_row, text="View Global Disqualifiers", command=self.show_disqualifier_reference).pack(side="right")

        # Notes section.
        notes_frame = ttk.Frame(frm)
        notes_frame.pack(fill="both", expand=True)

        ttk.Label(notes_frame, text="Question notes").pack(anchor="w")
        q_text = tk.Text(notes_frame, height=6, wrap="word")
        q_text.pack(fill="x", pady=4)
        q_text.insert(END, state.get("question_notes", ""))

        ttk.Label(notes_frame, text="Trait notes").pack(anchor="w")
        t_text = tk.Text(notes_frame, height=6, wrap="word")
        t_text.pack(fill="x", pady=4)
        t_text.insert(END, state.get("trait_notes", ""))

        ttk.Label(notes_frame, text="Verbatim quote/notes (required if disqualifier checked)").pack(anchor="w")
        v_text = tk.Text(notes_frame, height=6, wrap="word")
        v_text.pack(fill="x", pady=4)
        v_text.insert(END, state.get("verbatim_notes", ""))

        def persist_state() -> bool:
            """
            Save current screen inputs back into InterviewState.
            Returns False if a required field is missing.
            """
            raw = raw_var.get()
            if raw not in {1, 2, 3, 4, 5}:
                messagebox.showerror("Validation", "Raw score is required (select 1-5).")
                return False

            self.state.trait_inputs[tid]["raw_score"] = raw
            self.state.trait_inputs[tid]["absolute_disqualifier"] = bool(dq_var.get())
            self.state.trait_inputs[tid]["question_notes"] = q_text.get("1.0", END).strip()
            self.state.trait_inputs[tid]["trait_notes"] = t_text.get("1.0", END).strip()
            self.state.trait_inputs[tid]["verbatim_notes"] = v_text.get("1.0", END).strip()

            # Track current position (1-based) for drafts.
            self.state.current_index = idx + 1
            return True

        def go_back():
            if not persist_state():
                return
            self.show_trait_screen(idx - 1)

        def go_next():
            if not persist_state():
                return
            if idx + 1 < len(self.active_traits):
                self.show_trait_screen(idx + 1)
            else:
                # End of traits; you could show a summary screen.
                messagebox.showinfo("End", "You reached the final trait. Use Finalize to export.")
                self.scroll_top()

        def save_draft():
            if not persist_state():
                return
            try:
                dm = DraftManager(Path(self.settings["base_dir"]))
                payload = self.state.to_dict()
                path = dm.save_draft(payload)
                messagebox.showinfo("Draft Saved", f"Draft saved to:\n{path}")
            except Exception as exc:
                messagebox.showerror("Draft Save Error", f"{exc}\n\n{traceback.format_exc()}")

        def finalize():
            if not persist_state():
                return
            try:
                self.validate_before_finalize()
                scoring = ScoringEngine.evaluate(self.rubric, self.state.track, self.state.trait_inputs)
                payload = self.state.to_dict()

                exporter = DocxExporter(Path(self.settings["base_dir"]) / "final")
                out_path = exporter.export(self.rubric, payload, scoring)

                messagebox.showinfo(
                    "Finalized",
                    (
                        f"Outcome: {scoring['outcome']}\n"
                        f"Weighted Total: {scoring['weighted_total']}/{scoring['max_weighted_total']}\n"
                        f"Percent: {scoring['percent_of_max']}%\n\n"
                        f"Report saved to:\n{out_path}"
                    ),
                )
                self.show_start_screen()
            except Exception as exc:
                messagebox.showerror("Finalize Error", f"{exc}\n\n{traceback.format_exc()}")

        self.set_footer_actions(
            left_actions=[
                ("Back", go_back),
                ("Next", go_next),
                ("Save Draft", save_draft),
            ],
            right_actions=[
                ("Finalize", finalize),
                ("Exit", self.destroy),
            ],
        )

    def validate_before_finalize(self) -> None:
        """
        Final validation gate before exporting:
        - candidate name, interview date, school, track present
        - every trait has a valid score
        - disqualifier checked requires verbatim evidence
        """
        if not self.state.candidate_name.strip():
            raise ValueError("Candidate Name is required.")
        if not is_valid_date_yyyy_mm_dd(self.state.interview_date.strip()):
            raise ValueError("Interview Date must be valid YYYY-MM-DD.")
        if not self.state.school.strip():
            raise ValueError("School selection is required.")
        if not self.state.track:
            raise ValueError("Track selection is required.")

        traits = self.rubric_loader.get_traits_for_track(self.state.track)
        for trait in traits:
            tid = trait["id"]
            tstate = self.state.trait_inputs.get(tid)
            if not tstate or tstate.get("raw_score") not in {1, 2, 3, 4, 5}:
                raise ValueError(f"Missing raw score for trait: {trait['name']}")

            # Your rule: if disqualifier is checked, verbatim notes must be present.
            if tstate.get("absolute_disqualifier") and not (tstate.get("verbatim_notes") or "").strip():
                raise ValueError(f"Trait '{trait['name']}' has disqualifier checked but no verbatim notes.")


# =========================
# Entrypoint
# =========================

if __name__ == "__main__":
    try:
        app = InterviewApp()
        app.mainloop()
    except Exception as exc:
        # Fallback fatal error dialog.
        # Note: messagebox requires a Tk root; if Tk fails to init, this may not show.
        messagebox.showerror("Fatal Error", f"{exc}\n\n{traceback.format_exc()}")
